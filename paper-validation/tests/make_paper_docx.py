"""선택된 두 논문에 대한 '아주 자세한' 한국어 설명 docx 생성.

각 논문 폴더(papers/<folder>/) 에 paper_detailed_KR.docx 를 만든다.
내용은 PDF 본문에서 추출한 실제 수치·데이터·방법론에 근거한다.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PAPERS = ROOT / "papers"

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)


# ── 렌더링 헬퍼 ──
def add_title(doc: Document, text: str, sub: str) -> None:
    h = doc.add_heading(text, level=0)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    p = doc.add_paragraph()
    r = p.add_run(sub)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.text = k
        c1.text = v
        for r in c0.paragraphs[0].runs:
            r.bold = True
        c0.width = Pt(120)


def h1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = ACCENT


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Number")


def result_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light List Accent 1"
    for i, htext in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = htext
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


# ════════════════════════════════════════════════════════════════════
# 논문 #3: Diffusion-VAE (D-Va)
# ════════════════════════════════════════════════════════════════════
def build_dva() -> None:
    doc = Document()
    add_title(
        doc,
        "Diffusion Variational Autoencoder (D-Va)",
        "Tackling Stochasticity in Multi-Step Regression Stock Price Prediction",
    )

    meta_table(
        doc,
        [
            ("저자", "Kelvin J.L. Koa, Yunshan Ma, Ritchie Ng, Tat-Seng Chua (NUS / Eastspring)"),
            ("발표", "CIKM '23 (ACM Int. Conf. on Information and Knowledge Management)"),
            ("연도 / 인용", "2023 / 약 40회"),
            ("arXiv", "2309.00073  [q-fin.ST]"),
            ("코드", "https://github.com/koa-fin/dva (공개)"),
            ("분류", "딥러닝 주가예측 · 생성모델(VAE+Diffusion)"),
            ("재현 난이도", "중 (코드 공개 + 무료 주가데이터로 재현 가능)"),
        ],
    )

    h1(doc, "1. 한 줄 요약")
    para(
        doc,
        "주가의 강한 확률성(stochasticity)을 다루기 위해 계층적 VAE와 확산(diffusion) "
        "기법을 결합하여, 단일 시점이 아닌 '여러 날의 수익률 시퀀스'를 회귀 예측하는 "
        "seq2seq 생성모델 D-Va를 제안한다. 입력·타깃 양쪽에 노이즈를 주입(diffusion)하고 "
        "예측을 디노이징(denoising score-matching)해, 데이터 잡음(aleatoric)과 모델 "
        "불확실성(epistemic)을 함께 줄인다.",
    )

    h1(doc, "2. 문제 정의 (Problem Formulation)")
    para(
        doc,
        "각 종목 s에 대해 과거 T 거래일의 입력 시퀀스 X = {x_{t-T+1}, ..., x_t} 가 주어졌을 때, "
        "다음 T' 거래일의 수익률 시퀀스 y = {r_{t+1}, ..., r_{t+T'}} 를 예측한다. "
        "여기서 r_t = c_t / c_{t-1} (c = 종가)이다.",
    )
    bullets(
        doc,
        [
            "입력 피처 x_t = [open, high, low, volume, 절대수익률 Δ_t, 퍼센트수익률 r_t] (다변량 입력).",
            "open/high/low 는 전일 종가로 정규화 (예: o_t = o_t / c_{t-1}).",
            "다변량 입력 → 단변량(수익률) 출력, 멀티스텝 회귀 과제.",
            "동기: 멀티스텝 예측은 변동성 예측·파생상품 헤지·규제 유동성 기간(최소 10일) 대응에 필요.",
        ],
    )

    h1(doc, "3. 방법론 — 4개 핵심 구성요소")
    h2(doc, "(1) 계층적 VAE (Hierarchical VAE)")
    para(
        doc,
        "이미지 생성용 SOTA 모델인 NVAE(Nouveau VAE)를 seq2seq 예측용으로 재설계. "
        "depthwise separable convolution + batch norm + Swish 활성화 + Squeeze-and-Excitation(SE) "
        "게이팅을 사용한 인코더/디코더 residual cell로 구성. 잠재변수 스택 Z={Z1,Z2,Z3}을 통해 "
        "주가의 복잡한 저수준 잠재요인을 표현. 예측 시퀀스의 데이터 밀도: p(ŷ_n|X_n) = ∫ p_θ(Z|X_n)·δ(ŷ_n − f(Z)) dZ.",
    )
    h2(doc, "(2) 입력 시퀀스 확산 (X-Diffusion)")
    para(
        doc,
        "마르코프 연쇄로 입력 X에 점진적 가우시안 노이즈를 추가. 분산 스케줄 β_n∈[0,1]. "
        "reparameterization으로 임의 단계 n의 샘플을 닫힌형으로 추출: "
        "X_n = √(ᾱ_n)·X + √(1−ᾱ_n)·ε, ε~N(0,I), ᾱ_n=∏(1−β_i). "
        "→ 입력에 단계적 증강을 가해 잡음에 강건한 학습 유도.",
    )
    h2(doc, "(3) 타깃 시퀀스 결합 확산 (Y-Diffusion, coupled)")
    para(
        doc,
        "타깃 y에도 결합 노이즈를 주입(β'_n = γ·β_n). 생성분포와 확산분포의 KL을 맞추면 "
        "(식 4의 부등식) 생성모델이 만드는 전체 불확실성이 감소함을 이용. "
        "손실 L_KL = D_KL( p(ŷ_n) || q(y_n) ).",
    )
    h2(doc, "(4) 디노이징 스코어 매칭 (Denoising Score-Matching)")
    para(
        doc,
        "표준 확산모델의 역과정을 예측기로 대체. 에너지함수 기울기 ∇E(ŷ)를 학습해, "
        "테스트 시 한 번의 디노이징 점프 ŷ_final = ŷ − ∇E(ŷ) 로 잡음 ε_y(aleatoric)를 제거. "
        "타깃을 '정확히' 복원하지 않고 실제 매니폴드 y_r에 가까운 저차원 매니폴드를 학습.",
    )
    h2(doc, "최종 손실함수")
    para(doc, "L = L_MSE + ζ·L_KL + η·L_DSM   (그리드서치로 ζ=0.5, η=1 채택)")

    h1(doc, "4. 데이터셋")
    result_table(
        doc,
        ["데이터셋(테스트연도)", "기간", "종목수", "거래일수"],
        [
            ["2016", "2014-01-01 ~ 2016-12-31", "88", "756"],
            ["2019", "2017-01-01 ~ 2019-12-31", "110", "754"],
            ["2022", "2020-01-01 ~ 2022-12-31", "110", "756"],
        ],
    )
    bullets(
        doc,
        [
            "기본: ACL18 StockNet 데이터셋(미국 고거래량 88종목, 9개 산업 상위 8~10종목).",
            "확장: 2017~2023 데이터를 Yahoo Finance에서 수집(11개 산업 상위 10종목 → 110종목).",
            "각 3년 데이터셋을 train:val:test = 7:1:2 로 시간순 분할.",
            "시퀀스 길이 T = T' = 10, 20, 40, 60일 (규제 유동성 기간에 대응).",
        ],
    )

    h1(doc, "5. 실험 설정")
    bullets(
        doc,
        [
            "베이스라인: ARIMA, NBA(LSTM+attention, 텍스트 제거), VAE, VAE+Adversarial, Autoformer.",
            "최적화: Adam, 초기 lr 5e-4, 배치 16, 종목별 20 epoch, 종목마다 5회 반복.",
            "평가지표: 평균 MSE 및 표준편차(SD), 최강 베이스라인 대비 개선율(%).",
        ],
    )

    h1(doc, "6. 주요 결과 (논문 수치 = 검증 대상)")
    h2(doc, "6.1 예측 정확도 (MSE) — 발췌")
    result_table(
        doc,
        ["연도/T", "Autoformer", "D-Va", "MSE 개선", "SD 개선"],
        [
            ["2016 / 10", "1.0204", "0.9040", "11.41%", "73.06%"],
            ["2019 / 10", "1.1011", "0.9847", "10.57%", "77.87%"],
            ["2022 / 10", "0.9958", "0.8653", "13.10%", "70.23%"],
            ["2022 / 60", "0.8428", "0.8174", "3.01%", "78.22%"],
        ],
    )
    para(
        doc,
        "→ D-Va는 최강 베이스라인 대비 평균 MSE 7.49% 개선, 예측 표준편차(불확실성) "
        "평균 75.01% 감소. 예측 길이 T가 길수록 MSE 개선폭은 감소(장기엔 예기치 못한 충격 증가).",
    )
    h2(doc, "6.2 포트폴리오 (RQ3) — 10일 Sharpe Ratio")
    para(
        doc,
        "예측 수익률의 평균·공분산으로 Markowitz 평균-분산 최적화(무공매도 제약) 수행, "
        "공분산은 graphical lasso(λ=0.1)로 정규화. 무위험금리 0 가정.",
    )
    result_table(
        doc,
        ["테스트연도", "NBA", "Equal-Weight", "D-Va", "D-Va (정규화)"],
        [
            ["2016", "0.0270", "0.1089", "0.0772", "0.1174"],
            ["2019", "0.0820", "0.2337", "0.1197", "0.2767"],
            ["2022", "0.0332", "0.0437", "0.0600", "0.0645"],
        ],
    )
    para(
        doc,
        "→ D-Va 예측 + 공분산 정규화 조합이 모든 연도에서 최고 Sharpe. 정규화 없는 D-Va는 "
        "강력 베이스라인인 동일가중을 항상 이기진 못함(뉴스 등 외부정보 미반영 때문).",
    )

    h1(doc, "7. 핵심 클레임 (재현 시 직접 비교할 수치)")
    bullets(
        doc,
        [
            "C1: 최강 베이스라인 대비 평균 MSE 7.49% 개선.",
            "C2: 예측 표준편차 평균 75.01% 감소(불확실성 저감).",
            "C3: ablation — 각 구성요소(VAE→Xd→Yd→Dn)가 불확실성 높은 구간에서 MSE/SD 개선.",
            "C4: D-Va+정규화 포트폴리오가 NBA·동일가중 대비 최고 10일 Sharpe.",
        ],
    )

    h1(doc, "8. 우리 무료 셋업으로의 재현 계획")
    h2(doc, "데이터 매핑")
    bullets(
        doc,
        [
            "주가: Tiingo adjClose 사용(논문은 Yahoo). OHLCV+수익률 피처 동일 구성 가능.",
            "주의: 논문은 분할/배당 조정 정도가 다를 수 있음 → Tiingo adjClose 기준임을 리포트에 명시.",
            "공개 코드(koa-fin/dva)를 기반으로 데이터 로더만 우리 data_loader로 교체.",
        ],
    )
    h2(doc, "검증 4단계")
    numbered(
        doc,
        [
            "In-sample 재현: 논문과 동일 유니버스(미국 대형주)·2016/2019/2022 구간으로 MSE·SD·Sharpe 재현.",
            "OOS 종목 확장: 다른 섹터/지수 구성종목 ≥3개 그룹으로 일반화 검증.",
            "Time robustness: 2023~2025 등 신규 구간 ≥3개로 안정성 검증.",
            "비교표+시각화: 논문 수치 vs 재현 vs 확장(누적수익·Sharpe·MSE) 표와 Plotly 차트.",
        ],
    )

    h1(doc, "9. 리스크 / 주의사항")
    bullets(
        doc,
        [
            "종목별 개별 학습(20 epoch×110종목×5run) → 계산량 큼. GPU 권장, 소규모 유니버스부터 시작.",
            "Yahoo↔Tiingo 가격 차이(생존편향·조정방식)로 In-sample 수치가 미세하게 달라질 수 있음(caveat).",
            "포트폴리오 Sharpe는 graphical lasso 등 후처리에 민감 → 동일 하이퍼파라미터(λ=0.1) 사용.",
        ],
    )

    out = PAPERS / "2023_DiffusionVAE_Koa" / "paper_detailed_KR.docx"
    doc.save(str(out))
    print("saved:", out)


# ════════════════════════════════════════════════════════════════════
# 논문 #8: Transformer-DRL + Black-Litterman (BDA)
# ════════════════════════════════════════════════════════════════════
def build_bda() -> None:
    doc = Document()
    add_title(
        doc,
        "Transformer 기반 DRL + Black-Litterman (BDA)",
        "Combining Transformer based Deep Reinforcement Learning with the Black-Litterman Model for Portfolio Optimization",
    )

    meta_table(
        doc,
        [
            ("저자", "Ruoyu Sun, Angelos Stefanidis, Zhengyong Jiang, Jionglong Su (XJTLU)"),
            ("연도 / 인용", "2024 / 약 25회"),
            ("arXiv", "2402.16609  [q-fin.CP / cs.LG]"),
            ("분류", "강화학습 포트폴리오 최적화 · Black-Litterman · Transformer"),
            ("재현 난이도", "중상 (주가는 무료로 가능, BL+DRL 구현 필요)"),
        ],
    )

    h1(doc, "1. 한 줄 요약")
    para(
        doc,
        "일반적인 DRL 포트폴리오 에이전트는 자산 간 '동적 상관관계'를 학습하지 못해 "
        "(특히 공매도가 가능한 미국시장에서) 위험 대비 수익을 키우기 어렵다. 이 논문은 "
        "DRL 에이전트가 베이지안 기법인 Black-Litterman(BL) 모델을 '적용하는 정책'을 "
        "학습하게 하여, 자산 상관을 반영한 롱/숏 전략을 구사하는 BDA(Black-Litterman model "
        "based DRL Agent)를 제안한다.",
    )

    h1(doc, "2. 문제의식 / 동기")
    bullets(
        doc,
        [
            "전통 DRL 정책망은 보상(수익-위험)만으로 학습 → 자산 간 동적 상관을 정책에 담기 어려움.",
            "Markowitz 평균-분산은 추정오차 최대화(error maximization) 문제 → OOS 성능 저하.",
            "BL 모델은 베이지안으로 '주관적 view + 과거기반 prior'를 결합해 추정오차 문제를 완화.",
            "→ DRL이 BL 모델의 입력(기대수익 view, 위험회피)을 신경망으로 출력하도록 학습.",
        ],
    )

    h1(doc, "3. 방법론 — BDA 구조")
    bullets(
        doc,
        [
            "DRL 에이전트가 'BL 모델을 적용하는 정책'을 학습 → 목표 포트폴리오 가중치 산출.",
            "변형 Transformer(τ1) + CNN(τ2)이 BL의 주관적 기대수익 view와 위험회피를 출력.",
            "Transformer에서 position encoding 제거 → 여러 자산 수익률 시계열 간 '비선형 상관' "
            "학습에 집중, 과적합 완화·일반화 향상.",
            "prior 분포 파라미터는 과거 데이터로 계산.",
            "학습: Jiang et al.(EIIE)식 결정론적 정책경사(DPG). critic의 차원의 저주를 피하려 "
            "환경 보상함수로 목적함수를 직접 구성하고 해석적 기울기를 정책망에 역전파.",
            "보상함수: 수익 − 위험 − 거래규모(거래비용·차입비용 포함), 공매도 허용·레버리지 제한(총 투자=초기자본의 절반).",
        ],
    )
    para(
        doc,
        "정의: 가격벡터 v_t = 각 트레이딩 기간 마지막 거래일의 adjusted close 벡터. "
        "트레이딩 기간 = 5거래일(주간 리밸런싱). 무위험자산(현금) 가중치 w_{0,t}=1−Σw_{i,t}. "
        "거래비용율 c, 현금/주식 차입율 포함. 총 학습 스텝 3e5.",
    )

    h1(doc, "4. 데이터셋")
    bullets(
        doc,
        [
            "유니버스: 다우존스(DJIA) 구성종목, 결측 제거 후 29종목.",
            "출처: Yahoo Finance, adjusted close(배당·분할 조정).",
            "리밸런싱: 5거래일(주간), 거래비용·차입비용 반영.",
        ],
    )
    h2(doc, "4개 롤링 실험 (train 3년 → backtest 120거래일)")
    result_table(
        doc,
        ["실험", "학습 기간", "백테스트 시작"],
        [
            ["Exp1", "2018-01-01 ~ 2020-12-31", "2021-01-01부터 120거래일"],
            ["Exp2", "2018-07-01 ~ 2021-06-30", "2021-07-01부터 120거래일"],
            ["Exp3", "2019-01-01 ~ 2021-12-31", "2022-01-01부터 120거래일"],
            ["Exp4", "2019-07-01 ~ 2022-06-30", "2022-07-01부터 120거래일"],
        ],
    )

    h1(doc, "5. 실험 설정")
    bullets(
        doc,
        [
            "베이스라인(다수): CRP, UBAH, UP, EG, ANTICOR, PAMR, CWMR, OLMAR, ONS, BK 등 "
            "온라인 포트폴리오 선택 전략 + 여러 DRL 프레임워크.",
            "평가지표: AR(누적수익), DR(일수익), Std, LStd(하방 표준편차), SR(Sharpe), STR(Sortino).",
            "거래비용·차입비용을 매수/매도에 모두 포함.",
        ],
    )

    h1(doc, "6. 주요 결과 (논문 수치 = 검증 대상)")
    h2(doc, "6.1 실험 1 정량 결과 (발췌)")
    result_table(
        doc,
        ["전략", "AR(누적수익)", "DR", "Std", "SR", "STR"],
        [
            ["BDA", "0.3971", "0.00331", "0.02374", "0.1394", "0.1390"],
            ["ONS", "0.2079", "0.00173", "0.01066", "0.1624", "0.1662"],
            ["UP", "0.2001", "0.00167", "0.01092", "0.1527", "0.1526"],
            ["CRP", "0.1999", "0.00167", "0.01092", "0.1526", "0.1525"],
            ["Anticor", "0.1736", "0.00145", "0.01323", "0.1093", "0.1069"],
            ["OLMAR", "-0.0594", "-0.00049", "0.02400", "-0.0206", "-0.0199"],
            ["PAMR", "-0.2274", "-0.00189", "0.02250", "-", "-"],
        ],
    )
    para(
        doc,
        "→ 실험1에서 BDA 누적수익 39.7%로 차선(ONS 20.8%) 대비 약 +91%. SR은 실험1에서 ONS가 "
        "근소 우위지만, 실험 2·3·4에서는 BDA가 SR·STR 최고. BDA의 SR·STR은 전 실험에서 0.08 "
        "초과(다른 전략은 달성 못함).",
    )
    h2(doc, "6.2 헤드라인 클레임")
    bullets(
        doc,
        [
            "C1: 누적수익에서 비교전략·대체 DRL 대비 '최소 42%' 우위.",
            "C2: 4개 실험 모두에서 최고 누적수익(AR).",
            "C3: 실험 2·3·4에서 최고 Sharpe·Sortino, 전 실험 SR/STR > 0.08.",
            "C4: ablation — 보상함수 기반 목적함수 직접 역전파가 과적합 완화·일반화에 기여.",
        ],
    )

    h1(doc, "7. 우리 무료 셋업으로의 재현 계획")
    h2(doc, "데이터 매핑")
    bullets(
        doc,
        [
            "주가: Tiingo adjClose로 DJIA 29종목 수집(논문은 Yahoo). 주간(5거래일) 리밸런싱 동일.",
            "무위험금리: FRED(예: DGS1MO/DGS3MO)로 차입·현금금리 근사.",
            "거래비용·차입비용·레버리지 제한(총투자=초기자본 절반) 동일 적용.",
        ],
    )
    h2(doc, "검증 4단계")
    numbered(
        doc,
        [
            "In-sample 재현: DJIA 29종목·4개 롤링 구간으로 AR·SR·STR 재현(특히 BDA AR vs 베이스라인).",
            "OOS 종목 확장: S&P100/섹터 ETF 등 다른 유니버스 ≥3개로 일반화 검증.",
            "Time robustness: 2023~2025 등 신규 백테스트 구간 ≥3개로 안정성 검증.",
            "비교표+시각화: 논문 AR/SR vs 재현 vs 확장 표 + 누적수익 곡선·위험-수익 산점도.",
        ],
    )

    h1(doc, "8. 리스크 / 주의사항")
    bullets(
        doc,
        [
            "공개 코드 명시 없음 → BL 모델 + DRL(결정론적 정책경사) + 변형 Transformer를 직접 구현해야 함(구현난이도 중상).",
            "공매도·레버리지·거래비용 가정이 결과에 민감 → 논문 설정을 그대로 재현해야 비교 유효.",
            "Yahoo↔Tiingo 가격 조정 차이로 In-sample 수치가 미세하게 달라질 수 있음(caveat 명시).",
            "DJIA 구성종목은 시점에 따라 변동 → 백테스트 구간의 실제 구성종목/생존편향 주의.",
        ],
    )

    out = PAPERS / "2024_TransformerDRL_BlackLitterman_Sun" / "paper_detailed_KR.docx"
    doc.save(str(out))
    print("saved:", out)


if __name__ == "__main__":
    build_dva()
    build_bda()
