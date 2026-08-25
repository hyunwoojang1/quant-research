"""선택된 두 논문을 '중학생도 이해할 수 있게' 설명하는 쉬운 한국어 docx 생성.

기존 paper_detailed_KR.docx 를 덮어쓴다. 모든 어려운 개념은 일상 비유로 풀고,
끝에 용어 사전을 둔다. 사실(수치·데이터)은 PDF 본문에 근거.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PAPERS = ROOT / "papers"
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
GREEN = RGBColor(0x1E, 0x6B, 0x3A)


def add_title(doc, text, sub):
    h = doc.add_heading(text, level=0)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    p = doc.add_paragraph()
    r = p.add_run(sub)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = ACCENT


def h2(doc, text):
    doc.add_heading(text, level=2)


def para(doc, text):
    doc.add_paragraph(text)


def easy(doc, text):
    """'쉽게 말하면' 비유 박스 (Intense Quote 스타일 + 초록 라벨)."""
    p = doc.add_paragraph(style="Intense Quote")
    lab = p.add_run("쉽게 말하면  ")
    lab.bold = True
    lab.font.color.rgb = GREEN
    p.add_run(text)


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def meta_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.text, c1.text = k, v
        for r in c0.paragraphs[0].runs:
            r.bold = True


def table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light List Accent 1"
    for i, htext in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = htext
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def glossary(doc, items):
    h1(doc, "용어 사전 (어려운 말 쉽게 풀이)")
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = "용어"
    t.rows[0].cells[1].text = "쉬운 뜻"
    for r in t.rows[0].cells:
        for run in r.paragraphs[0].runs:
            run.bold = True
    for term, mean in items:
        cells = t.add_row().cells
        cells[0].text = term
        cells[1].text = mean
        for run in cells[0].paragraphs[0].runs:
            run.bold = True


COMMON_GLOSSARY = [
    ("주식 / 종목", "회사를 잘게 나눈 '조각 소유권'. 그 조각의 가격이 주가."),
    ("수익률", "어제보다 오늘 가격이 몇 % 올랐는지(또는 내렸는지)."),
    ("변동성", "가격이 얼마나 심하게 출렁이는지. 클수록 위험."),
    ("포트폴리오", "여러 주식에 돈을 나눠 담은 '바구니'. 한 곳에 몰빵 안 하기."),
    ("Sharpe(샤프) 지수", "'위험을 감수한 것에 비해 얼마나 잘 벌었나' 점수. 높을수록 좋음."),
    ("Sortino(소르티노) 지수", "샤프와 비슷하지만 '손실 위험'만 따져 계산한 점수."),
    ("백테스트", "과거 데이터로 '만약 그때 이 전략을 썼다면?' 모의실험."),
    ("In-sample(인샘플) 재현", "논문이 쓴 바로 그 종목·그 기간으로 결과를 그대로 따라 해보기."),
    ("신경망(딥러닝)", "사람 뇌의 신경세포를 흉내 낸, 데이터로 스스로 배우는 컴퓨터 모델."),
    ("MSE(오차)", "예측이 정답에서 얼마나 벗어났는지. 작을수록 정확."),
]


# ════════════════════════════════════════════════════════════════════
# 논문 #3: Diffusion-VAE (D-Va) — 쉬운 버전
# ════════════════════════════════════════════════════════════════════
def build_dva():
    doc = Document()
    add_title(
        doc,
        "주가를 여러 날 미리 맞히는 AI: Diffusion-VAE (D-Va)",
        "어려운 제목: Diffusion Variational Autoencoder for Multi-Step Stock Price Prediction (Koa 외, 2023)",
    )

    meta_table(
        doc,
        [
            ("누가 썼나", "싱가포르국립대(NUS) 연구팀 + 투자회사 Eastspring"),
            ("언제 / 얼마나 유명", "2023년 / 다른 논문이 약 40번 인용(꽤 유명)"),
            ("원문 주소", "arXiv 2309.00073"),
            ("코드 공개?", "네! github.com/koa-fin/dva 에서 받을 수 있음"),
            ("한마디로", "주가를 '내일 하루'가 아니라 '앞으로 여러 날' 미리 맞히는 AI"),
        ],
    )

    h1(doc, "1. 이 논문이 풀려는 문제")
    para(
        doc,
        "주가는 예측하기 정말 어렵습니다. 이유가 두 가지 있어요. 첫째, 주가는 럭비공처럼 "
        "어디로 튈지 모르게 마구 움직입니다(=무작위성이 큼). 둘째, 우리가 보는 '하루 종가'는 "
        "하루 종일 출렁인 가격 중 딱 한 순간만 찍은 사진 같아서, 진짜 흐름을 다 담지 못해요.",
    )
    para(
        doc,
        "게다가 대부분의 기존 연구는 '내일 오를까 내릴까?' 하루짜리만 맞혔습니다. 하지만 "
        "큰 금융회사나 은행은 '앞으로 여러 날'의 흐름을 알아야 위험을 관리할 수 있어요.",
    )
    easy(
        doc,
        "내일 날씨만 맞히는 게 아니라, 앞으로 열흘~두 달 날씨를 한꺼번에 맞히려는 것과 같아요. "
        "훨씬 어렵지만 훨씬 쓸모 있죠.",
    )

    h1(doc, "2. 어떻게 풀었나 (핵심 아이디어 3가지)")

    h2(doc, "아이디어 1 — VAE: '핵심만 요약했다가 다시 그리는 화가'")
    para(
        doc,
        "VAE는 복잡한 그림을 보고 핵심 특징만 짧게 메모해뒀다가, 그 메모만으로 그림을 다시 "
        "그려내는 화가 같은 모델입니다. 이 논문은 주가의 '숨은 핵심 패턴'을 이렇게 압축해서 배웁니다.",
    )

    h2(doc, "아이디어 2 — Diffusion(확산): '일부러 모래 뿌렸다 닦기 연습'")
    para(
        doc,
        "깨끗한 사진에 모래(노이즈)를 조금씩 뿌려 점점 흐릿하게 만든 뒤, 다시 그 모래를 "
        "걷어내 원래 사진을 복원하는 연습을 반복합니다. 이렇게 '지저분한 데이터 다루기'를 "
        "미리 훈련해두면, 잡음 많은 진짜 주가 앞에서도 흔들리지 않게 됩니다.",
    )
    easy(
        doc,
        "시험 공부할 때 일부러 어려운 문제·함정 문제를 많이 풀어보면, 실제 시험에서 어떤 문제가 "
        "나와도 안 당황하는 것과 같아요.",
    )

    h2(doc, "아이디어 3 — 정답에도 모래 뿌렸다가 마지막에 깨끗이 닦기")
    para(
        doc,
        "보통은 입력(과거 데이터)에만 노이즈를 넣지만, 이 논문은 정답(미래 수익률)에도 노이즈를 "
        "넣어 함께 훈련합니다. 그리고 예측이 끝나면 마지막에 한 번에 '쓱' 닦아내(denoising) "
        "더 깔끔하고 일반적인 예측을 만듭니다.",
    )

    h1(doc, "3. 어떤 데이터로 실험했나")
    table(
        doc,
        ["실험 묶음", "기간", "종목 수"],
        [
            ["2016년 테스트", "2014~2016년", "미국 주식 88개"],
            ["2019년 테스트", "2017~2019년", "미국 주식 110개"],
            ["2022년 테스트", "2020~2022년", "미국 주식 110개"],
        ],
    )
    bullets(
        doc,
        [
            "미국의 거래량 많은 대표 주식들(여러 산업에서 골고루)을 사용.",
            "데이터를 '공부용 : 연습용 : 시험용 = 7 : 1 : 2'로 나눔(시간 순서대로).",
            "10일, 20일, 40일, 60일 등 여러 기간을 미리 맞히도록 실험.",
        ],
    )

    h1(doc, "4. 결과 — 얼마나 잘했나")
    para(doc, "기존의 가장 좋은 방법들과 비교했을 때:")
    bullets(
        doc,
        [
            "예측이 틀린 정도(오차)가 평균 약 7.5% 줄었어요 → 더 정확.",
            "예측이 들쭉날쭉한 정도(불안정함)가 약 75% 줄었어요 → 훨씬 일관됨.",
        ],
    )
    easy(
        doc,
        "같은 시험을 5번 봐도 점수가 거의 안 흔들리고(일관됨), 평균 점수도 더 높아진(정확함) 셈이에요.",
    )
    h2(doc, "투자에 써보니? (포트폴리오 점수)")
    para(
        doc,
        "예측 결과로 '어느 주식에 얼마씩 담을지' 정해서 투자 바구니를 만들어 봤더니, "
        "위험 대비 수익 점수(Sharpe)가 가장 좋았습니다. 숫자가 클수록 좋아요.",
    )
    table(
        doc,
        ["테스트 연도", "단순 똑같이 나눠담기", "이 논문(D-Va)"],
        [
            ["2016", "0.109", "0.117"],
            ["2019", "0.234", "0.277"],
            ["2022", "0.044", "0.065"],
        ],
    )

    h1(doc, "5. 우리가 직접 따라 해볼 계획 (무료 데이터로)")
    numbered(
        doc,
        [
            "그대로 따라하기: 논문과 같은 미국 주식·같은 기간으로 결과가 재현되는지 확인.",
            "다른 주식으로: 논문에 없던 다른 종목·업종 3개 이상으로도 통하는지 시험.",
            "다른 시기로: 2023~2025년 같은 새 기간 3개 이상에서도 잘 되는지 시험.",
            "표로 비교: 논문 숫자 vs 우리 숫자를 나란히 표로 정리.",
        ],
    )
    para(
        doc,
        "데이터는 무료인 Tiingo(주가)와 FRED(금리)를 씁니다. 논문은 야후 파이낸스를 썼는데, "
        "데이터 출처가 조금 달라 숫자가 미세하게 다를 수 있어요(이 점은 보고서에 표시합니다). "
        "다행히 코드가 공개돼 있어 따라 하기 비교적 쉽습니다.",
    )

    h1(doc, "6. 주의할 점")
    bullets(
        doc,
        [
            "주식마다 따로 학습해서 컴퓨터 계산이 많이 필요(가능하면 GPU, 적은 종목부터 시작).",
            "뉴스 같은 외부 정보는 안 쓰기 때문에, 갑작스러운 사건은 못 맞힐 수 있음.",
        ],
    )

    glossary(
        doc,
        COMMON_GLOSSARY
        + [
            ("VAE", "복잡한 데이터를 핵심만 압축했다가 다시 만들어내는 AI(요약+복원 화가)."),
            ("Diffusion(확산) 모델", "노이즈를 일부러 더했다 지우는 연습으로 배우는 AI."),
            ("멀티스텝 예측", "하루가 아니라 여러 날을 한꺼번에 미리 맞히기."),
        ],
    )

    out = PAPERS / "2023_DiffusionVAE_Koa" / "paper_detailed_KR.docx"
    doc.save(str(out))
    print("saved:", out)


# ════════════════════════════════════════════════════════════════════
# 논문 #8: Transformer-DRL + Black-Litterman (BDA) — 쉬운 버전
# ════════════════════════════════════════════════════════════════════
def build_bda():
    doc = Document()
    add_title(
        doc,
        "스스로 배우는 투자 로봇: Transformer DRL + Black-Litterman (BDA)",
        "어려운 제목: Combining Transformer based Deep RL with the Black-Litterman Model for Portfolio Optimization (Sun 외, 2024)",
    )

    meta_table(
        doc,
        [
            ("누가 썼나", "시안교통-리버풀대(XJTLU) 연구팀"),
            ("언제 / 얼마나 유명", "2024년 / 다른 논문이 약 25번 인용"),
            ("원문 주소", "arXiv 2402.16609"),
            ("한마디로", "게임하듯 스스로 배우는 AI가, 여러 주식에 돈을 똑똑하게 나눠 담는 법"),
        ],
    )

    h1(doc, "1. 이 논문이 풀려는 문제")
    para(
        doc,
        "'강화학습'은 게임을 하면서 점수(보상)를 받아 스스로 잘하는 법을 깨치는 AI예요. "
        "이걸로 투자 로봇을 만들어 '여러 주식에 돈을 어떻게 나눠 담을지' 시킬 수 있습니다.",
    )
    para(
        doc,
        "그런데 기존 투자 로봇에는 약점이 있어요. 'A 주식이 오르면 B 주식은 내린다' 같은 "
        "주식들 사이의 관계(상관관계)를 잘 못 배웁니다. 이 관계를 알아야 한쪽이 손해 볼 때 "
        "다른 쪽이 메워주는 안전한 투자를 할 수 있는데 말이죠.",
    )
    easy(
        doc,
        "축구팀을 짤 때 공격수만 잔뜩 뽑으면 안 되고, 공격·수비가 서로 보완되게 뽑아야 하죠. "
        "기존 로봇은 이 '서로 보완하는 관계'를 잘 못 봤어요.",
    )

    h1(doc, "2. 어떻게 풀었나 (핵심 아이디어)")
    h2(doc, "비결 1 — Black-Litterman: '내 생각'과 '시장 통계'를 섞는 믹서기")
    para(
        doc,
        "Black-Litterman(BL)은 ① 투자자의 예상(내 생각)과 ② 과거 데이터에서 나온 시장의 "
        "통계적 사실을 적절히 섞어 균형 잡힌 결론을 내주는 도구입니다. 한쪽으로 치우치지 "
        "않게 막아줘요.",
    )
    para(
        doc,
        "이 논문의 핵심 아이디어: 로봇이 가중치를 막 정하게 두지 않고, 'BL 믹서기에 넣을 내 "
        "의견'만 정하도록 학습시킵니다. 그러면 믹서기가 알아서 균형을 잡아줘 더 안정적이에요.",
    )
    easy(
        doc,
        "요리할 때 재료를 아무렇게나 넣지 않고, 정해진 좋은 레시피(BL)에 '내 입맛 조절'만 "
        "더하는 것과 같아요. 망칠 위험이 줄죠.",
    )
    h2(doc, "비결 2 — Transformer: 여러 주식을 '동시에' 읽는 똑똑한 눈")
    para(
        doc,
        "Transformer는 요즘 챗GPT에도 쓰이는 신경망입니다. 이걸로 여러 주식의 움직임을 한꺼번에 "
        "읽어 '어떤 주식이 좋아 보인다'는 의견을 만듭니다. 특별히 '순서 정보'는 일부러 빼서, "
        "순서보다 '주식들끼리의 관계'에 더 집중하게 했어요(과한 암기를 막는 효과).",
    )
    h2(doc, "비결 3 — 공매도(롱/숏)까지 활용")
    para(
        doc,
        "이 로봇은 오를 주식은 사고(롱), 내릴 주식은 미리 팔아(숏) 양쪽으로 돈을 법니다. "
        "거래 수수료와 빌리는 비용도 계산에 넣어 현실적으로 실험했어요.",
    )

    h1(doc, "3. 어떤 데이터로 실험했나")
    bullets(
        doc,
        [
            "미국 대표 30개 기업 지수(다우존스, DJIA) 종목 중 29개 사용.",
            "한 주(5거래일)마다 한 번씩 투자 비중을 다시 조정(주간 리밸런싱).",
            "2018~2022년을 4번에 나눠, 매번 3년 공부 → 그 다음 120일로 실전 모의투자.",
        ],
    )
    table(
        doc,
        ["회차", "공부 기간", "모의투자 시작"],
        [
            ["1회", "2018.01~2020.12", "2021.01부터"],
            ["2회", "2018.07~2021.06", "2021.07부터"],
            ["3회", "2019.01~2021.12", "2022.01부터"],
            ["4회", "2019.07~2022.06", "2022.07부터"],
        ],
    )

    h1(doc, "4. 결과 — 얼마나 잘했나")
    para(doc, "1회차 모의투자 결과를 예로 보면(누적수익 = 기간 동안 총 몇 % 벌었나):")
    table(
        doc,
        ["전략", "누적수익", "위험대비수익(샤프)"],
        [
            ["이 논문(BDA)", "+39.7%", "0.139"],
            ["기존 최고(ONS)", "+20.8%", "0.162"],
            ["보통 전략(CRP)", "+20.0%", "0.153"],
            ["나쁜 예(PAMR)", "-22.7%", "(손실)"],
        ],
    )
    bullets(
        doc,
        [
            "4번의 모의투자 모두에서 누적수익 1등.",
            "기존 어떤 전략보다도 최소 42% 더 많이 벌었음.",
            "위험 대비 수익(샤프)도 2·3·4회차에서 1등(1회차만 근소한 2등).",
        ],
    )
    easy(
        doc,
        "여러 번 시험을 봤는데 '얼마나 벌었나'에서 매번 1등을 했고, '위험을 생각해도 잘했나'에서도 "
        "대부분 1등이었다는 뜻이에요.",
    )

    h1(doc, "5. 우리가 직접 따라 해볼 계획 (무료 데이터로)")
    numbered(
        doc,
        [
            "그대로 따라하기: 같은 다우존스 29종목·같은 4개 기간으로 수익·샤프가 재현되는지 확인.",
            "다른 주식으로: S&P100이나 업종 ETF 등 다른 묶음 3개 이상으로도 통하는지 시험.",
            "다른 시기로: 2023~2025년 같은 새 기간 3개 이상에서도 잘 되는지 시험.",
            "표로 비교: 논문 숫자 vs 우리 숫자를 나란히 표로 정리.",
        ],
    )
    para(
        doc,
        "주가는 무료 Tiingo, 금리는 무료 FRED를 씁니다. 단, 이 논문은 코드가 공개돼 있지 "
        "않아서 'BL 믹서기 + 강화학습 로봇 + Transformer'를 우리가 직접 만들어야 합니다 "
        "(앞 논문보다 손이 더 많이 감).",
    )

    h1(doc, "6. 주의할 점")
    bullets(
        doc,
        [
            "공개 코드가 없어 직접 구현해야 함 → 시간이 더 걸림.",
            "공매도·수수료·빌리는 비용 설정에 결과가 민감 → 논문과 똑같이 맞춰야 공정한 비교.",
            "다우존스 구성 종목은 시기마다 바뀜 → 그 시점의 실제 종목으로 맞춰야 함.",
        ],
    )

    glossary(
        doc,
        COMMON_GLOSSARY
        + [
            ("강화학습(DRL)", "게임처럼 점수(보상)를 받으며 스스로 잘하는 법을 배우는 AI."),
            ("Transformer", "여러 정보를 동시에 보고 관계를 파악하는 신경망(챗GPT에도 쓰임)."),
            ("Black-Litterman", "'내 예상'과 '시장 통계'를 섞어 균형 잡힌 답을 내는 투자 도구."),
            ("공매도(숏)", "내릴 것 같은 주식을 미리 팔아, 값이 내리면 이익 보는 방법."),
            ("리밸런싱", "투자 비중을 주기적으로 다시 조정하는 것."),
            ("누적수익", "기간 전체 동안 합쳐서 몇 % 벌었는지."),
        ],
    )

    out = PAPERS / "2024_TransformerDRL_BlackLitterman_Sun" / "paper_detailed_KR.docx"
    doc.save(str(out))
    print("saved:", out)


if __name__ == "__main__":
    build_dva()
    build_bda()
