"""선택된 두 논문에 대한 '전문 리포트 수준' 한국어 설명 docx 생성 (v3).

기존 paper_detailed_KR.docx 를 덮어쓴다.
- 톤: 대학원 세미나 리딩노트 수준(전문적). 유치한 비유 대신 '직관적 해석' 보조설명.
- 전문용어는 처음 등장 시 정의.
- 연구 전체 워크플로우 포함: 모델 선택 근거 → 데이터 전처리 → 아키텍처 → 학습 →
  평가 프로토콜 → 결과/ablation → 재현 워크플로우 → 한계 → 용어집.
- 사실(수식·수치·데이터)은 PDF 본문에 근거.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PAPERS = ROOT / "papers"
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
MUTE = RGBColor(0x44, 0x44, 0x44)
INTUIT = RGBColor(0x6A, 0x4A, 0x00)


def add_title(doc, kr, en):
    h = doc.add_heading(kr, level=0)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    p = doc.add_paragraph()
    r = p.add_run(en)
    r.italic = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTE


def h1(doc, t):
    h = doc.add_heading(t, level=1)
    for r in h.runs:
        r.font.color.rgb = ACCENT


def h2(doc, t):
    doc.add_heading(t, level=2)


def p(doc, text):
    doc.add_paragraph(text)


def intuition(doc, text):
    """전문 보조설명('직관적 해석'). 유치하지 않게, 본문 흐름을 돕는 한 줄 해석."""
    par = doc.add_paragraph()
    lab = par.add_run("직관적 해석 — ")
    lab.bold = True
    lab.italic = True
    lab.font.color.rgb = INTUIT
    run = par.add_run(text)
    run.italic = True
    run.font.color.rgb = MUTE


def eq(doc, text):
    """수식 줄: 들여쓰기 + 고정폭 느낌."""
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Pt(18)
    r = par.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def meta_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.text, c1.text = k, v
        for r in c0.paragraphs[0].runs:
            r.bold = True


def table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light List Accent 1"
    for i, ht in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ht
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v


def concept_table(doc, items):
    """전문용어 정의 표 (정의 + 한 줄 직관)."""
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    for i, ht in enumerate(["용어", "정의", "직관"]):
        t.rows[0].cells[i].text = ht
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    for term, defi, intu in items:
        cells = t.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = term, defi, intu
        for r in cells[0].paragraphs[0].runs:
            r.bold = True


# ════════════════════════════════════════════════════════════════════
# 논문 #3: Diffusion-VAE (D-Va)
# ════════════════════════════════════════════════════════════════════
def build_dva():
    doc = Document()
    add_title(
        doc,
        "D-Va: 확산 변분오토인코더 기반 다중스텝 주가 예측",
        "Diffusion Variational Autoencoder for Tackling Stochasticity in Multi-Step "
        "Regression Stock Price Prediction (Koa, Ma, Ng, Chua — CIKM '23)",
    )
    meta_table(
        doc,
        [
            ("저자/소속", "Kelvin J.L. Koa, Yunshan Ma, Tat-Seng Chua (NUS); Ritchie Ng (Eastspring Investments)"),
            ("발표/연도", "ACM CIKM 2023 / 인용 약 40회"),
            ("arXiv / 분류", "2309.00073 / q-fin.ST (computational finance)"),
            ("코드", "공개 — https://github.com/koa-fin/dva"),
            ("과제 유형", "다변량 입력 → 단변량(수익률) 멀티스텝 회귀 (seq2seq 생성모델)"),
            ("핵심 기여", "계층적 VAE + 입력/타깃 결합확산 + 디노이징으로 데이터·모델 불확실성 동시 처리"),
        ],
    )

    h1(doc, "1. 연구 배경과 문제 정의")
    p(
        doc,
        "주가 예측이 어려운 본질적 이유는 두 가지다. (a) 주가는 강한 확률성(stochasticity)을 "
        "가져 표준 예측모델이 잘 일반화되지 않는다. (b) 우리가 관측하는 일별 종가는 연속적으로 "
        "변동하는 가격을 특정 시점에 이산적으로 '표집'한 값이라, 가격의 본질적 거동을 온전히 "
        "담지 못한다(관측 잡음).",
    )
    p(
        doc,
        "또한 기존 연구의 대부분은 '다음 1스텝 상승/하락'의 이진 분류에 집중했고, 표현력이 낮은 "
        "모델에 머물렀다. 그러나 변동성 예측·파생상품 헤지·규제상 유동성 기간(최소 10일) 대응을 "
        "위해서는 '여러 날의 수익률 시퀀스'를 예측하는 멀티스텝 회귀가 필요하다. 멀티스텝에서는 "
        "타깃 시퀀스 자체에도 잡음이 섞여 학습 시 일반화가 더 어려워진다.",
    )
    p(
        doc,
        "정식 정의: 종목 s에 대해 과거 T 거래일 입력 X = {x_{t−T+1}, …, x_t} 로부터 향후 "
        "T′ 거래일의 수익률 시퀀스 y = {r_{t+1}, …, r_{t+T′}} 를 예측한다 (r_t = c_t / c_{t−1}).",
    )

    h1(doc, "2. 사전 개념 정리 (전문용어 정의)")
    concept_table(
        doc,
        [
            ("VAE (변분오토인코더)",
             "입력을 확률적 잠재변수 z의 분포로 인코딩하고 다시 복원하도록 학습하는 생성모델. "
             "재구성오차 + KL규제(잠재분포를 사전분포에 근접)로 학습.",
             "데이터의 핵심을 확률적으로 압축·복원하는 모델"),
            ("계층적 VAE / NVAE",
             "잠재변수를 여러 층(Z1,Z2,Z3)으로 쌓아 표현력을 키운 VAE. NVAE는 이미지 생성용 SOTA로, "
             "depthwise separable conv·BatchNorm·Swish·SE 게이팅을 사용.",
             "단일 잠재변수보다 복잡한 패턴을 더 잘 포착"),
            ("확산모델 (Diffusion/DDPM)",
             "데이터에 가우시안 노이즈를 점진적으로 더하는 정과정과, 이를 되돌리는 역과정을 "
             "학습하는 생성모델.",
             "노이즈 주입·제거를 반복 학습해 잡음에 강해짐"),
            ("디노이징 스코어매칭(DSM)",
             "노이즈가 섞인 샘플에서 깨끗한 데이터 매니폴드로 향하는 스코어(로그밀도 기울기)를 "
             "학습하는 기법.",
             "예측을 실제 데이터 면(manifold)으로 끌어당김"),
            ("Aleatoric / Epistemic 불확실성",
             "전자는 데이터 자체의 무작위성, 후자는 모델·추정의 불확실성.",
             "잡음(데이터) vs 모름(모델)"),
        ],
    )

    h1(doc, "3. 모델 선택의 근거 (왜 이 구성인가)")
    bullets(
        doc,
        [
            "왜 회귀(분류 아님)? — 포트폴리오 비중 산정을 위해 수익률의 크기/순위가 필요. "
            "이진 분류로는 자산 가중을 못 한다.",
            "왜 계층적 VAE? — 단일 잠재변수 VAE는 표현력이 낮음. 계층적 잠재변수 Z={Z1,Z2,Z3}로 "
            "주가의 복잡·저수준 잠재요인을 더 풍부히 표현(이미지 SOTA인 NVAE를 seq2seq로 전용).",
            "왜 확산(입력+타깃)? — 주가의 강한 확률성을 '노이즈 주입 학습'으로 모사하여 일반화·"
            "강건성 확보. 타깃에도 결합확산을 적용하면(식 4의 KL 부등식) 생성 불확실성이 감소.",
            "왜 디노이징? — 타깃 y가 본질적으로 잡음(y = y_r + ε_y)이므로, 한 번의 디노이징 점프로 "
            "aleatoric 잡음을 제거해 실제 매니폴드 y_r에 근접.",
            "왜 graphical lasso? — 포트폴리오 단계에서 공분산 추정의 모델(epistemic) 불확실성을 "
            "L1 규제로 축소(공분산 수축과 유사).",
        ],
    )
    intuition(
        doc,
        "데이터 잡음(aleatoric)은 확산+디노이징으로, 모델 불확실성(epistemic)은 공분산 정규화로 "
        "각각 분담해 처리하는 것이 이 논문의 설계 철학이다.",
    )

    h1(doc, "4. 데이터: 수집 → 전처리 파이프라인")
    h2(doc, "4.1 데이터셋")
    table(
        doc,
        ["데이터셋(테스트연도)", "기간", "종목수", "거래일수"],
        [
            ["2016", "2014-01-01 ~ 2016-12-31", "88", "756"],
            ["2019", "2017-01-01 ~ 2019-12-31", "110", "754"],
            ["2022", "2020-01-01 ~ 2022-12-31", "110", "756"],
        ],
    )
    bullets(
        doc,
        [
            "기본: ACL18 StockNet(미국 고거래량 88종목, 9개 산업 상위 8~10종목, 2014~2017).",
            "확장: 2017~2023 데이터를 Yahoo Finance에서 수집(11개 산업 상위 10종목 → 110종목), "
            "ACL18과 동일 방식으로 처리. 길이 일관성을 위해 3년×3개 데이터셋으로 분할.",
        ],
    )
    h2(doc, "4.2 피처 구성과 정규화 (전처리 핵심)")
    p(doc, "각 시점 입력 벡터는 6개 피처로 구성된다:")
    eq(doc, "x_t = [ o_t, h_t, l_t, v_t, Δ_t, r_t ]   (open, high, low, volume, 절대수익, 퍼센트수익)")
    bullets(
        doc,
        [
            "퍼센트수익률 r_t = c_t / c_{t−1}, 절대수익 Δ_t = c_t − c_{t−1}.",
            "open/high/low는 전일 종가로 정규화: o_t ← o_t / c_{t−1} (스케일 제거, 수익률 단위로 통일).",
            "다변량 입력(6피처) → 단변량 출력(수익률 시퀀스).",
            "윈도잉: 입력 길이 T와 출력 길이 T′ 를 {10, 20, 40, 60}으로 변화(규제 유동성 기간 반영).",
            "분할: 시간순 train:val:test = 7:1:2.",
        ],
    )

    h1(doc, "5. 모델 아키텍처 상세")
    h2(doc, "5.1 백본 — 계층적 VAE(NVAE) seq2seq")
    p(
        doc,
        "인코더 residual cell: (BatchNorm→Swish→Conv)×2 + Squeeze-and-Excitation(SE) 게이팅. "
        "디코더 residual cell: 위 구성 + depthwise separable convolution(수용영역 확대·연산량 절감). "
        "잠재변수 스택 Z={Z1,Z2,Z3}로 저수준 의존성 포착. 예측밀도:",
    )
    eq(doc, "p(ŷ_n | X_n) = ∫ p_θ(Z | X_n) · δ(ŷ_n − f(Z)) dZ")
    h2(doc, "5.2 입력 확산 (X-Diffusion)")
    eq(doc, "q(X_n|X)=N(X_n; √ᾱ_n·X, (1−ᾱ_n)I),  X_n=√ᾱ_n·X+√(1−ᾱ_n)·ε,  ε~N(0,I),  ᾱ_n=∏(1−β_i)")
    h2(doc, "5.3 타깃 결합 확산 (Y-Diffusion)")
    p(doc, "타깃에 결합 노이즈(β′_n=γ·β_n) 주입. 생성·확산 분포의 KL을 맞추면 전체 불확실성 감소:")
    eq(doc, "L_KL = D_KL( p(ŷ_n) || q(y_n) )")
    h2(doc, "5.4 디노이징 스코어매칭 (테스트 시 1-step jump)")
    eq(doc, "L_DSM,n = E[ σ_n · || y − ŷ_n + ∇_{ŷ_n} E(ŷ_n) ||² ] ,   ŷ_final = ŷ − ∇_{ŷ} E(ŷ)")

    h1(doc, "6. 학습 절차")
    p(doc, "총 손실 (트레이드오프 계수 ζ, η):")
    eq(doc, "L = L_MSE + ζ·L_KL + η·L_DSM        (그리드서치 → ζ=0.5, η=1)")
    bullets(
        doc,
        [
            "절차: (1) X,y에 결합확산 적용 → X_n,y_n 생성 (2) 계층적 VAE가 X_n→ŷ_n 생성, y_n에 매칭 "
            "(3) 동시에 디노이징 에너지함수 E(ŷ) 학습. 추론 시 1-step 디노이징으로 ŷ_final 산출.",
            "최적화: Adam, 초기 lr 5e-4, 배치 16, 종목별 20 epoch, 종목마다 5회 반복(평균 보고).",
        ],
    )

    h1(doc, "7. 평가 프로토콜")
    bullets(
        doc,
        [
            "지표: 종목·5회 평균 MSE 및 표준편차(SD), 최강 베이스라인 대비 개선율(%).",
            "베이스라인: ARIMA / NBA(LSTM+attention, 텍스트 제거) / VAE / VAE+Adversarial / Autoformer.",
            "포트폴리오(RQ3): 예측 평균·공분산으로 Markowitz 평균-분산 최적화(무공매도), "
            "공분산은 graphical lasso(λ=0.1) 정규화, 무위험금리 0, Sharpe로 평가.",
        ],
    )
    eq(doc, "포트폴리오:  max_w  wᵀμ − (γ/2)·wᵀΣw   s.t.  wᵀ1=1,  w≥0")

    h1(doc, "8. 결과 및 해석")
    h2(doc, "8.1 예측 정확도(MSE) — 발췌")
    table(
        doc,
        ["연도/T", "Autoformer", "D-Va", "MSE개선", "SD개선"],
        [
            ["2016 / 10", "1.0204", "0.9040", "11.41%", "73.06%"],
            ["2019 / 10", "1.1011", "0.9847", "10.57%", "77.87%"],
            ["2022 / 10", "0.9958", "0.8653", "13.10%", "70.23%"],
            ["2022 / 60", "0.8428", "0.8174", "3.01%", "78.22%"],
        ],
    )
    p(
        doc,
        "D-Va는 최강 베이스라인 대비 평균 MSE 7.49% 개선, 예측 SD 평균 75.01% 감소. "
        "예측 길이 T가 길수록 MSE 개선폭은 줄어드는데, 이는 장기 구간에서 예측 불가한 시장충격이 "
        "늘기 때문이다(저자는 짧은 모델의 롤링 예측을 대안으로 제시).",
    )
    h2(doc, "8.2 포트폴리오 — 10일 Sharpe")
    table(
        doc,
        ["테스트연도", "NBA", "Equal-Weight", "D-Va", "D-Va(정규화)"],
        [
            ["2016", "0.0270", "0.1089", "0.0772", "0.1174"],
            ["2019", "0.0820", "0.2337", "0.1197", "0.2767"],
            ["2022", "0.0332", "0.0437", "0.0600", "0.0645"],
        ],
    )

    h1(doc, "9. Ablation / 분석")
    bullets(
        doc,
        [
            "변형: D-Va−XdYdDn(백본 VAE만) / −YdDn / −Dn. 백본만으로도 ARIMA·Autoformer 상회.",
            "확산 구성요소 추가마다 예측 SD가 뚜렷이 감소(강건성↑).",
            "핵심 관찰: 직전 모델의 예측 불확실성(SD)이 높을수록 다음 구성요소의 MSE 개선이 큼 "
            "→ 구성요소들이 '잡음 처리'로 작동함을 시사.",
        ],
    )

    h1(doc, "10. 우리 환경 재현 워크플로우")
    bullets(
        doc,
        [
            "데이터 매핑: 주가 = Tiingo adjClose(논문은 Yahoo) · 금리 = FRED. OHLCV+수익률 피처 동일 구성.",
            "공개코드(koa-fin/dva) 기반, 데이터 로더만 shared/data_loader 로 교체.",
            "주의: Yahoo↔Tiingo 조정방식 차이로 in-sample 수치가 미세 변동 가능 → 리포트에 명시.",
        ],
    )
    numbered(
        doc,
        [
            "In-sample 재현: 동일 미국 대형주·2016/2019/2022 구간으로 MSE·SD·Sharpe 재현.",
            "OOS 종목 확장: 다른 섹터/지수 구성종목 ≥3개 그룹으로 일반화 검증.",
            "Time robustness: 2023~2025 등 신규 구간 ≥3개로 안정성 검증.",
            "비교표+Plotly 시각화: 논문 vs 재현 vs 확장(MSE·Sharpe·누적수익).",
        ],
    )

    h1(doc, "11. 한계와 재현 리스크")
    bullets(
        doc,
        [
            "종목별 개별 학습(20ep×110종목×5run) → 계산량 큼, GPU 권장·소규모부터.",
            "뉴스 등 외부정보 미사용 → 급격한 충격 예측 불가(비정규화 포트폴리오는 동일가중에 종종 열위).",
            "데이터 출처 차이(생존편향·조정)로 인한 수치 편차 가능.",
        ],
    )
    out = PAPERS / "2023_DiffusionVAE_Koa" / "paper_detailed_KR.docx"
    doc.save(str(out))
    print("saved:", out)


# ════════════════════════════════════════════════════════════════════
# 논문 #8: BDA (Transformer-DRL + Black-Litterman)
# ════════════════════════════════════════════════════════════════════
def build_bda():
    doc = Document()
    add_title(
        doc,
        "BDA: Black-Litterman 기반 Transformer 심층강화학습 포트폴리오 최적화",
        "Combining Transformer based Deep Reinforcement Learning with the "
        "Black-Litterman Model for Portfolio Optimization (Sun et al., 2024)",
    )
    meta_table(
        doc,
        [
            ("저자/소속", "Ruoyu Sun, Angelos Stefanidis, Zhengyong Jiang, Jionglong Su (XJTLU)"),
            ("연도", "2024 / 인용 약 25회"),
            ("arXiv / 분류", "2402.16609 / q-fin.CP, cs.LG"),
            ("과제 유형", "연속기간 포트폴리오 최적화(롱/숏 허용)를 위한 결정론적 정책경사 DRL"),
            ("핵심 기여", "DRL이 'BL 모델 입력(view·위험회피)'을 산출하는 정책을 학습 → 자산 간 동적 상관 반영"),
        ],
    )

    h1(doc, "1. 연구 배경과 문제 정의")
    p(
        doc,
        "심층강화학습(DRL)은 환경과 상호작용하며 보상으로 정책을 학습하는 모델-프리 기법으로, "
        "시장 변화에 동적으로 적응하고 자산 간 결합동학을 명시적으로 가정하지 않아도 된다는 "
        "장점이 있다. 그러나 통상의 DRL 포트폴리오 정책은 보상(수익·위험)만으로 학습되어 "
        "'자산 간 동적 상관관계'를 정책에 담기 어렵다. 상관관계는 롱/숏 전략의 핵심(음의 상관 "
        "자산이 서로 손익을 상쇄해 비체계적 위험을 줄임)이므로, 이를 못 쓰면 특히 공매도가 "
        "허용되는 미국시장에서 위험 대비 수익을 키우기 어렵다.",
    )
    p(
        doc,
        "한편 Markowitz 평균-분산은 추정오차 최대화(error maximization) 문제로 OOS 성능이 "
        "나빠진다. 이를 보완하는 Black-Litterman(BL) 모델은 베이지안으로 '주관적 view + "
        "과거기반 prior'를 결합해 안정적 해를 준다. 본 논문은 DRL이 BL 모델의 입력을 산출하도록 "
        "학습시키는 BDA(Black-Litterman model based DRL Agent)를 제안한다.",
    )

    h1(doc, "2. 사전 개념 정리 (전문용어 정의)")
    concept_table(
        doc,
        [
            ("강화학습(DRL)",
             "상태에서 행동을 취해 보상을 받고, 누적보상을 최대화하는 정책을 학습. 모델-프리.",
             "보상으로 시행착오하며 정책을 학습"),
            ("정책경사 / 결정론적 정책(DPG)",
             "정책을 직접 파라미터화해 목적함수의 기울기로 갱신. 결정론적 정책은 상태→행동을 "
             "확정적으로 매핑.",
             "행동확률 대신 행동 자체를 출력"),
            ("Black-Litterman(BL)",
             "시장균형(prior)과 투자자 view를 베이지안으로 결합해 사후 기대수익 분포를 산출하는 "
             "평균-분산 계열 모델.",
             "시장 통계와 내 의견을 신뢰도로 가중 결합"),
            ("Transformer",
             "self-attention으로 시퀀스 내 요소 간 관계를 병렬로 학습하는 신경망. 통상 위치인코딩으로 "
             "순서정보를 부여.",
             "여러 입력 간 상호관계를 동시에 포착"),
            ("EIIE / 정책-only 학습",
             "Jiang et al.의 포트폴리오 정책 토폴로지. 자산별 동일 평가기를 공유, critic 없이 보상 "
             "기반 목적함수를 직접 최적화.",
             "critic 없이 보상에서 바로 정책 학습"),
            ("롱/숏, 레버리지",
             "롱=매수, 숏=공매도. 레버리지는 원금 대비 큰 포지션. 본 논문은 총투자=초기자본의 0.5배로 제한.",
             "오를 건 사고 내릴 건 미리 팔되 규모 제한"),
        ],
    )

    h1(doc, "3. 모델 선택의 근거 (왜 이 구성인가)")
    bullets(
        doc,
        [
            "왜 BL(Markowitz 아님)? — Markowitz는 추정오차 최대화로 OOS가 불안정. BL은 균형 prior에 "
            "view를 신뢰도로 결합해 이를 완화하고, 평균-분산 구조라 상관 기반 롱/숏이 가능.",
            "왜 DRL이 'view·위험회피'만 출력? — 가중치를 직접 출력하면 행동공간이 고차원·연속이라 "
            "탐색이 비효율적. BL이라는 구조적 사전을 두고 그 입력만 학습하면 안정성↑.",
            "왜 Transformer + 위치인코딩 제거? — 여러 자산 수익률 시계열의 '비선형 상관'을 포착하되, "
            "위치(순서)정보를 제거해 시점 순서에 과적합되지 않고 자산 간 관계 학습에 집중.",
            "왜 정책-only(critic 제거)? — actor-critic은 고차원 연속행동에서 차원의 저주로 critic "
            "학습이 어려움. 보상함수로 미분가능 목적함수를 만들어 해석적 기울기를 정책망에 직접 역전파.",
        ],
    )

    h1(doc, "4. 데이터: 수집 → 전처리 파이프라인")
    bullets(
        doc,
        [
            "유니버스: DJIA 구성종목, 결측 제거 후 29종목. 출처 Yahoo Finance, adjusted close(배당·분할 조정).",
            "트레이딩 기간 = 5거래일(주간 리밸런싱). 기간 말 종가벡터로 거래 집행.",
            "로그수익률 사용: r_{i,t} = log2( c_{i,t} / c_{i,t−1} ).",
            "상태 텐서: 과거 m개 기간의 가격변동 텐서 X_t ∈ ℝ^{1×5m×N} (N=자산수).",
            "비용 모델: 수수료율 c, 현금·주식 차입율 반영. 무위험금리는 0 가정(2020~21 제로금리 반영).",
            "레버리지 제한: 각 기간 총투자액 p_t = 0.5·p_1 (롱+숏 규모 통제로 비용·위험 억제).",
        ],
    )
    p(doc, "상태 정의 (직전 비중 + 과거 수익 텐서):")
    eq(doc, "s_t := ( w_{t−1},  X_t ),   X_t = [ y_{t−m}, …, y_{t−1} ]")

    h1(doc, "5. 모델 아키텍처 & 의사결정 절차")
    h2(doc, "5.1 두 신경망의 역할")
    bullets(
        doc,
        [
            "τ1(Transformer, 위치인코딩 제거): 상태로부터 절대 view q_t(각 자산의 기대초과수익) 출력. "
            "각 기간 5일 수익을 패치로 분해해 시퀀스로 처리.",
            "τ2(CNN): 위험회피 계수 λ_t(스칼라) 출력.",
        ],
    )
    eq(doc, "q_t = τ1(s_t; θ1),   λ_t = τ2(s_t; θ2)")
    h2(doc, "5.2 BL 모델 계산 (prior → view → posterior → 가중치)")
    bullets(
        doc,
        [
            "prior 공분산 Σ_t^h: 과거 수익으로 추정(식 10).",
            "균형 prior 기대수익 Π_t: 역최적화로 도출(각 자산 동일 투자가치 가정, 식 11).",
            "view: DRL의 q_t를 절대 view로 사용 → view행렬 P=항등행렬(N×N), 불확실성 Ω=diag(τ·PΣP').",
            "사후분포(식 13)로 prior와 view를 결합해 μ_post, Σ_post 산출.",
        ],
    )
    eq(doc, "Σ_post,μ_post = BL( Π_t, Σ_t^h, P=I, Ω, q_t )      # 식(13)")
    p(doc, "최종 가중치는 오목 이차계획(QP)의 라그랑주 해(식 15):")
    eq(doc, "w_t = (1/λ_t)·Σ_post^{-1}·μ_post,   w_{0,t}=1 − Σ w_t   (현금)")

    h1(doc, "6. 학습 절차 (보상·목적함수)")
    p(doc, "보상함수 R: 일평균수익 − 분산패널티 − 거래규모패널티 (z1,z2>0):")
    eq(doc, "R(w_t | s_{t+1}, z1, z2) = (1/5)·μ_p − (z1/2)·σ_p² − z2·χ_p        # 식(17)")
    p(doc, "학습목표: 학습기간 누적보상 ARD 최대화. critic 없이 미분가능 목적함수의 해석적 기울기를 정책망에 역전파(정책-only, 차원의 저주 회피).")
    eq(doc, "ARD(tr) = Σ_t R(w_t | s_{t+1}, z1, z2)        # 식(18),  총 스텝 3e5")

    h1(doc, "7. 평가 프로토콜")
    bullets(
        doc,
        [
            "지표: AR(누적수익), DR(일수익), Std, LStd(하방표준편차), SR(Sharpe), STR(Sortino).",
            "베이스라인: CRP, UBAH, UP, EG, ANTICOR, PAMR, CWMR, OLMAR, ONS, BK 등 온라인 포트폴리오 "
            "전략 + 여러 DRL 프레임워크. 모든 매매에 거래·차입비용 포함.",
        ],
    )
    h2(doc, "롤링 백테스트 4구간 (train 3년 → backtest 120거래일)")
    table(
        doc,
        ["실험", "학습 기간", "백테스트 시작"],
        [
            ["Exp1", "2018-01-01 ~ 2020-12-31", "2021-01-01부터 120일"],
            ["Exp2", "2018-07-01 ~ 2021-06-30", "2021-07-01부터 120일"],
            ["Exp3", "2019-01-01 ~ 2021-12-31", "2022-01-01부터 120일"],
            ["Exp4", "2019-07-01 ~ 2022-06-30", "2022-07-01부터 120일"],
        ],
    )

    h1(doc, "8. 결과 및 해석")
    h2(doc, "8.1 실험 1 정량 결과 (발췌)")
    table(
        doc,
        ["전략", "AR(누적수익)", "DR", "Std", "SR", "STR"],
        [
            ["BDA", "0.3971", "0.00331", "0.02374", "0.1394", "0.1390"],
            ["ONS", "0.2079", "0.00173", "0.01066", "0.1624", "0.1662"],
            ["UP", "0.2001", "0.00167", "0.01092", "0.1527", "0.1526"],
            ["CRP", "0.1999", "0.00167", "0.01092", "0.1526", "0.1525"],
            ["Anticor", "0.1736", "0.00145", "0.01323", "0.1093", "0.1069"],
            ["OLMAR", "-0.0594", "-0.00049", "0.02400", "-0.0206", "-0.0199"],
        ],
    )
    p(
        doc,
        "Exp1에서 BDA 누적수익 39.7%로 차선(ONS 20.8%) 대비 약 +91%. SR은 Exp1에서 ONS가 근소 "
        "우위지만 Exp2·3·4에서는 BDA가 SR·STR 최고. BDA의 SR·STR은 전 실험에서 0.08 초과(타 전략은 미달). "
        "헤드라인 클레임: 누적수익에서 비교군·대체 DRL 대비 '최소 42%' 우위, 4개 실험 모두 AR 1위.",
    )

    h1(doc, "9. 우리 환경 재현 워크플로우")
    bullets(
        doc,
        [
            "데이터 매핑: 주가=Tiingo adjClose(DJIA 29종목, 논문은 Yahoo) · 무위험/차입금리=FRED(DGS1MO 등).",
            "주간(5거래일) 리밸런싱, 거래·차입비용·레버리지 0.5배 동일 적용.",
            "공개코드 없음 → BL + 결정론적 정책경사 + 위치인코딩 제거 Transformer를 직접 구현(난이도 중상).",
        ],
    )
    numbered(
        doc,
        [
            "In-sample 재현: DJIA 29종목·4개 롤링 구간으로 AR·SR·STR 재현(특히 BDA AR vs 베이스라인).",
            "OOS 종목 확장: S&P100/섹터 ETF 등 다른 유니버스 ≥3개로 일반화 검증.",
            "Time robustness: 2023~2025 등 신규 백테스트 구간 ≥3개로 안정성 검증.",
            "비교표+시각화: 논문 AR/SR vs 재현 vs 확장 + 누적수익 곡선·위험-수익 산점도.",
        ],
    )

    h1(doc, "10. 한계와 재현 리스크")
    bullets(
        doc,
        [
            "공개코드 없음 → 직접 구현 필요(BL QP·정책경사·Transformer). 검증의 최대 난관.",
            "공매도·수수료·차입·레버리지 가정에 결과가 민감 → 논문 설정을 정확히 복제해야 비교 유효.",
            "DJIA 구성종목은 시점별로 변동 → 백테스트 구간의 실제 구성·생존편향 주의.",
            "Yahoo↔Tiingo 조정 차이로 in-sample 수치 미세 편차 가능(caveat 명시).",
        ],
    )
    out = PAPERS / "2024_TransformerDRL_BlackLitterman_Sun" / "paper_detailed_KR.docx"
    doc.save(str(out))
    print("saved:", out)


if __name__ == "__main__":
    build_dva()
    build_bda()
